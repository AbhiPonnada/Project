from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
import pandas as pd
import google.generativeai as palm
from docx import Document
import os
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Load the Excel file
xls = pd.ExcelFile(r"data/data3.xlsx")

try:
    palm.configure(api_key="AIzaSyB-N4ZN2Hd7g3ftjPT23lgLUdUxph1-gkk")  # Replace with your actual API key
except Exception as e:
    print(f"Error initializing Google Generative AI: {e}")

dict ={
    "landscape": ['Component', 'Query', 'Query', '3.Landscape Details'],
    "materials": [ 'Type ', 'Question', 'Question', '5.Materials Management'],
    "production":['Type','Question','Question','6.Production Planning'],
    "quality":['Type ','Question','Question','7.Quality Management'],
    "warehouse":['Type','Question ','Question ','8.Warehouse Management'],
    "sales":['Type ','Question ','Question ','9.Sales Distribution'],
    "transportation":['Type ','Question ','Question ','10.Transportation Management'],
    "plant":['Type', 'Question', 'Question','11.Plant Maintenance'],
    "finance":['Process','Scope Question','Scope Question','12.Finance'],
    "control":['Sub-Process','Scope Question','Scope Question','13.Controlling,Treasury'],
    "project":['Type ','Question ','Question ','15.Project Systems  '],
    "start":['Category','Qualification Question','Qualification Question','1.Start Here'],
    "organisation":['Category','Questions','Questions','4.Organisation Structure'],
    "human":['Type ','Question ','Question ','14.Human Resources']
}

def get_values(input):
    [a,b,c,d] = dict[input]
    sheet_name = d
    df = pd.read_excel(xls, sheet_name)
    
    df['Category'] = df[a].fillna(method='ffill')
    questions_df = df.dropna(subset=[b])

    questions_by_category = {
        category: questions[c].tolist()
        for category, questions in questions_df.groupby('Category')
    }
    return a,b,c,d, questions_by_category

@app.route('/favicon.ico')
def favicon():
    return send_file('static/favicon.ico')

@app.route('/')
def website():
    return render_template('website.html')

@app.route('/<input>')
def index(input):
    a,b,c,d, questions_by_category = get_values(input)
    sheet_name = d
    
    title = sheet_name.split('.')[1]
    categories = list(questions_by_category.keys())
    return render_template('index.html', input = input, title = title, categories=categories)

@app.route('/<input>/<category>', methods=['GET', 'POST'])
def category(input, category):
    a,b,c,d, questions_by_category = get_values(input)
    session['user_answers'] = {cat: [""] * len(questions) for cat, questions in questions_by_category.items()}
    
    questions = questions_by_category[category]
    user_answers = session['user_answers'][category]
    suggestions = [""] * len(questions)

    if request.method == 'POST':
        for i, question in enumerate(questions):
            user_answers[i] = request.form.get(f'answer_{i}', '')
        
        session['user_answers'][category] = user_answers
        flash(f"Answers for {category} saved successfully!", 'success')
        return redirect(url_for('index', input= input))
    
    return render_template('category.html', input= input, category=category, questions=questions, user_answers=user_answers, suggestions=suggestions, enumerate=enumerate)

@app.route('/suggestion', methods=['POST'])
def suggestion():
    question = request.form['question']
    print(f"Received suggestion request for question: {question}")  # Log the received question
    suggestion = get_suggestion(question)
    print(f"Generated suggestion: {suggestion}")  # Log the generated suggestion
    return jsonify({'suggestion': suggestion})

@app.route('/<input>/summary', methods=['GET', 'POST'])
def summary(input):
    if 'user_answers' not in session:
        flash("No answers available to display in summary.", 'warning')
        return redirect(url_for('index', input= input))
    a,b,c,d, questions_by_category = get_values(input)
    summary_data = []
    if request.method == 'POST':
        user_name = request.form['user_name']
       
        for category, answers in session['user_answers'].items():
            questions = questions_by_category[category]
            for i, answer in enumerate(answers):
                question = questions[i]
                summary_data.append({'Category': category, 'Question': question, 'Answer': answer})
        
        summary_df = pd.DataFrame(summary_data)
        file_name = f"{user_name}_answers.xlsx"
        file_path = os.path.join('output', file_name)
        
        os.makedirs('output', exist_ok=True)  # Ensure the directory exists
        summary_df.to_excel(file_path, index=False)

        summary_text = "\n".join([f"{row['Question']}: {row['Answer']}" for _, row in summary_df.iterrows()])
        scope_suggestion = get_scope_suggestion(summary_text)
        word_file_name = f"{user_name}_scope_for_SAP.docx"
        word_file_path = os.path.join('output', word_file_name)
        
        doc = Document()
        doc.add_heading("Scope for SAP", level=1)
        doc.add_paragraph(scope_suggestion)
        doc.save(word_file_path)
        
        flash(f"Answers saved to {file_name} and Scope for SAP saved to {word_file_name}", 'success')
        return redirect(url_for('thank_you'))

    return render_template('summary.html', input=input, summary_data=summary_data, questions_by_category=questions_by_category, enumerate=enumerate)

@app.route('/thank_you')
def thank_you():
    return render_template('thank_you.html')

def get_suggestion(question):
    try:
        response = palm.generate_text(
            model='models/text-bison-001',
            prompt=f"Provide a suggestion for the following question according to Production Planning in SAP: {question}",
            max_output_tokens=150
        )
        return response.result
    except Exception as e:
        print(f"Error generating suggestion: {e}")  # Log the error
        return f"Error: {e}"

def get_scope_suggestion(summary_text):
    response = palm.generate_text(
        model='models/text-bison-001',
        prompt=f"Based on the following answers, what is the scope for SAP?\n\n{summary_text}",
        max_output_tokens=300
    )
    return response.result

@app.route('/download')
def download_output():
    output_folder = 'output'
    zip_file = 'results.zip'
    shutil.make_archive('results', 'zip', output_folder)
    return send_file(zip_file, as_attachment=True, download_name='results.zip')

if __name__ == '__main__':
    app.run(debug=True)
